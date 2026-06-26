#!/usr/bin/env python3
"""Build a styled DOCX from a simple Markdown report.

Supported Markdown:
  # / ## / ### headings
  paragraphs
  - bullets
  numbered lists starting with "1. "
  GitHub-style pipe tables
  image syntax: ![Alt](path/to/image.png)
  horizontal rule: ---

Usage:
  python scripts/build_docx.py report.md --output report.docx --font-east-asian 宋体
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
NUMBER_RE = re.compile(r"^\d+[.)]\s+(.*)$")


def set_run_font(run, east_asian: str, latin: str, size_pt: float | None = None, bold: bool | None = None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asian)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_style_font(style, east_asian: str, latin: str, size_pt: float | None = None, bold: bool | None = None):
    style.font.name = latin
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asian)
    if size_pt:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, east_asian: str, latin: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    set_run_font(run, east_asian, latin, size_pt=9, bold=bold)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def is_separator_line(line: str) -> bool:
    stripped = line.strip().strip("|").strip()
    if not stripped:
        return False
    return all(set(part.strip()) <= {"-", ":"} for part in stripped.split("|"))


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(document: Document, rows: list[list[str]], east_asian: str, latin: str):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text, east_asian, latin, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "F2F2F2")
    document.add_paragraph()


def add_paragraph_with_inline_bold(document: Document, text: str, style: str, east_asian: str, latin: str):
    p = document.add_paragraph(style=style)
    tokens = re.split(r"(\*\*.*?\*\*)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = p.add_run(token[2:-2])
            set_run_font(run, east_asian, latin, bold=True)
        else:
            run = p.add_run(token)
            set_run_font(run, east_asian, latin)
    return p


def configure_document(document: Document, east_asian: str, latin: str):
    section = document.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = document.styles
    set_style_font(styles["Normal"], east_asian, latin, 10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.3
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 16), ("Heading 2", 13.5), ("Heading 3", 12)]:
        set_style_font(styles[name], east_asian, latin, size, True)
        styles[name].paragraph_format.space_before = Pt(12)
        styles[name].paragraph_format.space_after = Pt(6)
    if "Title" in styles:
        set_style_font(styles["Title"], east_asian, latin, 20, True)


def add_page_break(document: Document):
    document.add_page_break()


def build_docx(markdown_path: Path, output_path: Path, east_asian: str, latin: str):
    document = Document()
    configure_document(document, east_asian, latin)
    lines = markdown_path.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            add_page_break(document)
            i += 1
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = heading.group(2).strip()
            if i == 0 and level == 1:
                p = document.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_run_font(run, east_asian, latin, 20, True)
            else:
                document.add_heading(text, level=level)
            i += 1
            continue

        if is_table_line(stripped):
            raw_rows = []
            while i < len(lines) and is_table_line(lines[i]):
                if not is_separator_line(lines[i]):
                    raw_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(document, raw_rows, east_asian, latin)
            continue

        image = IMAGE_RE.match(stripped)
        if image:
            alt, img_path_text = image.groups()
            img_path = Path(img_path_text)
            if not img_path.is_absolute():
                img_path = markdown_path.parent / img_path
            if img_path.exists():
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.4))
                if alt:
                    caption = document.add_paragraph(alt)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_paragraph_with_inline_bold(document, f"[Missing image: {img_path_text}]", "Normal", east_asian, latin)
            i += 1
            continue

        bullet = BULLET_RE.match(stripped)
        if bullet:
            p = add_paragraph_with_inline_bold(document, bullet.group(1), "List Bullet", east_asian, latin)
            i += 1
            continue

        numbered = NUMBER_RE.match(stripped)
        if numbered:
            p = add_paragraph_with_inline_bold(document, numbered.group(1), "List Number", east_asian, latin)
            i += 1
            continue

        add_paragraph_with_inline_bold(document, stripped, "Normal", east_asian, latin)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("--output", "-o", type=Path, default=Path("report.docx"))
    parser.add_argument("--font-east-asian", default="宋体")
    parser.add_argument("--font-latin", default="Times New Roman")
    args = parser.parse_args()
    build_docx(args.markdown, args.output, args.font_east_asian, args.font_latin)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
