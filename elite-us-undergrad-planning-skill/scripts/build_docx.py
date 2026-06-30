#!/usr/bin/env python3
"""Simple Markdown-to-DOCX builder for consulting-style Chinese reports.

This is intentionally conservative. It supports headings, paragraphs, bullets,
numbered lists, blockquotes, fenced code blocks, Markdown tables, and images.
For very complex documents, use this as a first pass and then inspect/polish.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


def load_style(path: Path | None) -> dict:
    if path and path.exists():
        return yaml.safe_load(path.read_text(encoding="utf-8"))
    return {
        "font": {"east_asia": "Microsoft YaHei", "latin": "Arial", "body_size_pt": 10.5},
        "page": {"margin_cm": 1.8},
        "colors": {"heading": "2F75B5", "table_header": "D9EAF7", "table_border": "BFBFBF"},
        "paragraph": {"line_spacing": 1.18, "space_after_pt": 4},
    }


def set_run_font(run, east_asia: str, latin: str, size_pt: float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "BFBFBF"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def apply_doc_style(doc: Document, style: dict):
    east = style["font"].get("east_asia", "Microsoft YaHei")
    latin = style["font"].get("latin", "Arial")
    body_size = style["font"].get("body_size_pt", 10.5)
    margin = Cm(float(style.get("page", {}).get("margin_cm", 1.8)))
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin
    normal = doc.styles["Normal"]
    normal.font.name = latin
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    normal.font.size = Pt(body_size)
    para_format = normal.paragraph_format
    para_format.line_spacing = style.get("paragraph", {}).get("line_spacing", 1.18)
    para_format.space_after = Pt(style.get("paragraph", {}).get("space_after_pt", 4))


def add_paragraph_with_inline(doc: Document, text: str, style_dict: dict, style_name: str | None = None):
    p = doc.add_paragraph(style=style_name)
    east = style_dict["font"].get("east_asia", "Microsoft YaHei")
    latin = style_dict["font"].get("latin", "Arial")
    # very simple bold inline support for **text**
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        run = p.add_run(clean)
        set_run_font(run, east, latin, bold=is_bold)
    return p


def is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def parse_table(lines: List[str], start: int):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and "|" in lines[i].strip()[1:]:
        table_lines.append(lines[i])
        i += 1
    if len(table_lines) < 2 or not is_table_separator(table_lines[1]):
        return None, start
    rows = []
    for j, line in enumerate(table_lines):
        if j == 1:
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: List[List[str]], style_dict: dict):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_fill = style_dict.get("colors", {}).get("table_header", "D9EAF7")
    border = style_dict.get("colors", {}).get("table_border", "BFBFBF")
    east = style_dict["font"].get("east_asia", "Microsoft YaHei")
    latin = style_dict["font"].get("latin", "Arial")
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, east, latin, size_pt=9.0, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
            set_cell_border(cell, border)
    doc.add_paragraph()


def add_image(doc: Document, md_path: Path, image_ref: str, caption: str | None = None):
    image_path = Path(image_ref)
    if not image_path.is_absolute():
        image_path = (md_path.parent / image_path).resolve()
    if not image_path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[图片缺失：{image_ref}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(6.4))
    except Exception:
        run.add_picture(str(image_path), width=Inches(5.8))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True


def build(md_path: Path, output: Path, style_path: Path | None):
    style_dict = load_style(style_path)
    doc = Document()
    apply_doc_style(doc, style_dict)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                if code_buf:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_buf))
                    run.font.name = "Consolas"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                    run.font.size = Pt(8.5)
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        table_rows, next_i = parse_table(lines, i)
        if table_rows is not None:
            add_table(doc, table_rows, style_dict)
            i = next_i
            continue

        img_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img_match:
            alt, ref = img_match.groups()
            add_image(doc, md_path, ref, alt)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_heading(text, level=min(level - 1, 4))
            for run in p.runs:
                set_run_font(run, style_dict["font"].get("east_asia", "Microsoft YaHei"), style_dict["font"].get("latin", "Arial"), bold=True, color=style_dict.get("colors", {}).get("heading", "2F75B5"))
            i += 1
            continue

        if stripped.startswith(">"):
            p = add_paragraph_with_inline(doc, stripped.lstrip("> "), style_dict)
            p.paragraph_format.left_indent = Cm(0.5)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            add_paragraph_with_inline(doc, text, style_dict, style_name="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+[.)]\s+", stripped):
            text = re.sub(r"^\d+[.)]\s+", "", stripped)
            add_paragraph_with_inline(doc, text, style_dict, style_name="List Number")
            i += 1
            continue

        add_paragraph_with_inline(doc, stripped, style_dict)
        i += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser(description="Build a DOCX from a Markdown report")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--style", type=Path, default=None)
    args = parser.parse_args()
    build(args.input, args.output, args.style)


if __name__ == "__main__":
    main()
