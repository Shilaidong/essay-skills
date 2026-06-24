#!/usr/bin/env python3
"""Build a polished DOCX from a YAML or JSON content file.

The schema is documented in assets/report-content-template.yaml.  The script
focuses on deterministic, editable Word output rather than complex desktop-
publishing effects.  It intentionally does not embed or distribute fonts.
"""

from __future__ import annotations

import argparse
import json
import sys
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, MutableMapping, Optional

try:
    import yaml
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency: PyYAML. Run: pip install PyYAML") from exc

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency: python-docx. Run: pip install python-docx") from exc


DEFAULT_STYLE: Dict[str, Any] = {
    "page_size": "A4",
    "east_asian_font": "宋体",
    "latin_font": "Times New Roman",
    "body_size_pt": 10.5,
    "line_spacing": 1.35,
    "heading_1_size_pt": 17,
    "heading_2_size_pt": 13.5,
    "heading_3_size_pt": 11.5,
    "heading_color": "#1F4E78",
    "accent_color": "#D9EAF7",
    "table_header_color": "#D9EAF7",
    "text_color": "#202020",
    "margin_top_cm": 2.3,
    "margin_bottom_cm": 2.2,
    "margin_left_cm": 2.4,
    "margin_right_cm": 2.4,
    "cover_page": True,
    "page_numbers": True,
    "header_text": "",
    "footer_text": "",
}


def deep_merge(base: MutableMapping[str, Any], override: Mapping[str, Any]) -> MutableMapping[str, Any]:
    """Recursively merge mapping values."""
    for key, value in override.items():
        if isinstance(value, Mapping) and isinstance(base.get(key), Mapping):
            base[key] = deep_merge(dict(base[key]), value)
        else:
            base[key] = value
    return base


def load_structured(path: Path) -> Dict[str, Any]:
    text = path.read_text(encoding="utf-8")
    if path.suffix.lower() == ".json":
        data = json.loads(text)
    else:
        data = yaml.safe_load(text)
    if not isinstance(data, dict):
        raise ValueError(f"Top-level content must be a mapping: {path}")
    return data


def rgb(hex_value: str) -> RGBColor:
    value = (hex_value or "#000000").strip().lstrip("#")
    if len(value) != 6:
        raise ValueError(f"Expected 6-digit hex color, got: {hex_value}")
    return RGBColor.from_string(value.upper())


def set_east_asian_font(element: Any, font_name: str) -> None:
    rpr = element.get_or_add_rPr() if hasattr(element, "get_or_add_rPr") else element
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_run_font(run: Any, style: Mapping[str, Any], *, size: Optional[float] = None,
                 bold: Optional[bool] = None, color: Optional[str] = None,
                 italic: Optional[bool] = None) -> None:
    run.font.name = str(style["latin_font"])
    set_east_asian_font(run._element, str(style["east_asian_font"]))
    if size is not None:
        run.font.size = Pt(float(size))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_style_font(doc_style: Any, style: Mapping[str, Any], size: float,
                   bold: bool = False, color: Optional[str] = None) -> None:
    doc_style.font.name = str(style["latin_font"])
    doc_style._element.rPr.rFonts.set(qn("w:eastAsia"), str(style["east_asian_font"]))
    doc_style.font.size = Pt(float(size))
    doc_style.font.bold = bold
    if color:
        doc_style.font.color.rgb = rgb(color)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.strip().lstrip("#").upper())


def set_cell_margins(cell: Any, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_text(paragraph: Any, text: str, style: Mapping[str, Any], *, bold: bool = False,
             italic: bool = False, size: Optional[float] = None, color: Optional[str] = None) -> Any:
    run = paragraph.add_run(str(text))
    set_run_font(run, style, size=size or float(style["body_size_pt"]), bold=bold,
                 italic=italic, color=color or str(style["text_color"]))
    return run


def configure_document(doc: Document, style: Mapping[str, Any], document_meta: Mapping[str, Any]) -> None:
    section = doc.sections[0]
    page_size = str(style.get("page_size", "A4")).upper()
    if page_size == "LETTER":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(float(style["margin_top_cm"]))
    section.bottom_margin = Cm(float(style["margin_bottom_cm"]))
    section.left_margin = Cm(float(style["margin_left_cm"]))
    section.right_margin = Cm(float(style["margin_right_cm"]))
    section.different_first_page_header_footer = bool(style.get("cover_page", True))

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, style, float(style["body_size_pt"]), color=str(style["text_color"]))
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = float(style["line_spacing"])

    for name, size in (
        ("Title", 24),
        ("Subtitle", 12),
        ("Heading 1", float(style["heading_1_size_pt"])),
        ("Heading 2", float(style["heading_2_size_pt"])),
        ("Heading 3", float(style["heading_3_size_pt"])),
    ):
        st = styles[name]
        set_style_font(st, style, size, bold=name != "Subtitle", color=str(style["heading_color"]))
        st.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            st.paragraph_format.space_before = Pt(10)
            st.paragraph_format.space_after = Pt(7)
        elif name in {"Heading 2", "Heading 3"}:
            st.paragraph_format.space_before = Pt(8)
            st.paragraph_format.space_after = Pt(4)

    for list_style_name in ("List Bullet", "List Number"):
        if list_style_name in styles:
            set_style_font(styles[list_style_name], style, float(style["body_size_pt"]), color=str(style["text_color"]))
            styles[list_style_name].paragraph_format.space_after = Pt(3)
            styles[list_style_name].paragraph_format.line_spacing = float(style["line_spacing"])

    if "Caption" in styles:
        set_style_font(styles["Caption"], style, 9, color="#666666")
        styles["Caption"].font.italic = False
        styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    core = doc.core_properties
    core.title = str(document_meta.get("title", ""))
    core.subject = str(document_meta.get("subtitle", ""))
    core.author = str(document_meta.get("author", ""))
    core.comments = "Generated with professional-document-studio"

    header_text = str(style.get("header_text", "") or "")
    footer_text = str(style.get("footer_text", "") or "")
    if header_text:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(hp, header_text, style, size=8.5, color="#666666")
    fp = section.footer.paragraphs[0]
    if footer_text:
        add_text(fp, footer_text + "   ", style, size=8.5, color="#666666")
    if bool(style.get("page_numbers", True)):
        add_page_number(fp)


def add_cover(doc: Document, meta: Mapping[str, Any], style: Mapping[str, Any]) -> None:
    for _ in range(3):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r = title_p.add_run(str(meta.get("title", "未命名文档")))
    set_run_font(title_r, style, size=24, bold=True, color=str(style["heading_color"]))

    subtitle = str(meta.get("subtitle", "") or "")
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, subtitle, style, size=12, color="#555555")

    doc.add_paragraph()
    details = [
        ("对象", meta.get("client", "")),
        ("作者", meta.get("author", "")),
        ("版本", meta.get("version", "")),
        ("日期", meta.get("date", "") or date.today().isoformat()),
        ("密级", meta.get("confidentiality", "")),
    ]
    for label, value in details:
        if value:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, f"{label}：{value}", style, size=10, color="#666666")

    doc.add_page_break()


def add_heading(doc: Document, title: str, level: int) -> None:
    level = max(1, min(3, int(level)))
    doc.add_heading(str(title), level=level)


def add_paragraph_block(doc: Document, block: Mapping[str, Any], style: Mapping[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = float(style["line_spacing"])
    align = str(block.get("align", "left")).lower()
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    add_text(
        p,
        str(block.get("text", "")),
        style,
        bold=bool(block.get("bold", False)),
        italic=bool(block.get("italic", False)),
        size=float(block.get("size_pt", style["body_size_pt"])),
        color=str(block.get("color", style["text_color"])),
    )


def add_list_block(doc: Document, block: Mapping[str, Any], style: Mapping[str, Any], numbered: bool) -> None:
    list_style = "List Number" if numbered else "List Bullet"
    for item in block.get("items", []) or []:
        if isinstance(item, Mapping):
            text = str(item.get("text", ""))
            level = int(item.get("level", 0))
        else:
            text = str(item)
            level = 0
        p = doc.add_paragraph(style=list_style)
        if level:
            p.paragraph_format.left_indent = Cm(0.6 * level)
        add_text(p, text, style)


def add_table_block(doc: Document, block: Mapping[str, Any], style: Mapping[str, Any]) -> None:
    headers = [str(x) for x in (block.get("headers") or [])]
    rows = block.get("rows") or []
    if not headers and not rows:
        return
    ncols = len(headers) if headers else max(len(row) for row in rows)
    table = doc.add_table(rows=1 if headers else 0, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    if headers:
        header_row = table.rows[0]
        repeat_table_header(header_row)
        for idx, value in enumerate(headers):
            cell = header_row.cells[idx]
            set_cell_shading(cell, str(style["table_header_color"]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, value, style, bold=True, size=9.5)

    for row_data in rows:
        row = table.add_row()
        for idx in range(ncols):
            value = row_data[idx] if idx < len(row_data) else ""
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.LEFT
            if isinstance(value, (int, float)):
                align = WD_ALIGN_PARAGRAPH.RIGHT
            p.alignment = align
            add_text(p, str(value), style, size=9.2)

    caption = str(block.get("caption", "") or "")
    if caption:
        p = doc.add_paragraph(style="Caption")
        add_text(p, caption, style, size=9, color="#666666")


def add_callout_block(doc: Document, block: Mapping[str, Any], style: Mapping[str, Any]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, str(block.get("fill", style["accent_color"])))
    set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
    p = cell.paragraphs[0]
    label = str(block.get("label", "重点"))
    text = str(block.get("text", ""))
    add_text(p, f"{label}｜", style, bold=True, color=str(style["heading_color"]))
    add_text(p, text, style)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_quote_block(doc: Document, block: Mapping[str, Any], style: Mapping[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    add_text(p, str(block.get("text", "")), style, italic=True, color="#555555")
    source = str(block.get("source", "") or "")
    if source:
        add_text(p, f"  — {source}", style, size=9, color="#777777")


def add_image_block(doc: Document, block: Mapping[str, Any], style: Mapping[str, Any], base_dir: Path) -> None:
    raw_path = str(block.get("path", ""))
    if not raw_path:
        return
    image_path = Path(raw_path)
    if not image_path.is_absolute():
        image_path = (base_dir / image_path).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"Image not found: {image_path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    width_cm = block.get("width_cm")
    height_cm = block.get("height_cm")
    kwargs: Dict[str, Any] = {}
    if width_cm is not None:
        kwargs["width"] = Cm(float(width_cm))
    if height_cm is not None:
        kwargs["height"] = Cm(float(height_cm))
    run.add_picture(str(image_path), **kwargs)
    caption = str(block.get("caption", "") or "")
    if caption:
        cp = doc.add_paragraph(style="Caption")
        add_text(cp, caption, style, size=9, color="#666666")


def add_blocks(doc: Document, blocks: Iterable[Mapping[str, Any]], style: Mapping[str, Any], base_dir: Path) -> None:
    for block in blocks:
        if not isinstance(block, Mapping):
            continue
        block_type = str(block.get("type", "paragraph")).lower()
        if block_type == "paragraph":
            add_paragraph_block(doc, block, style)
        elif block_type == "bullets":
            add_list_block(doc, block, style, numbered=False)
        elif block_type == "numbered":
            add_list_block(doc, block, style, numbered=True)
        elif block_type == "table":
            add_table_block(doc, block, style)
        elif block_type == "callout":
            add_callout_block(doc, block, style)
        elif block_type == "quote":
            add_quote_block(doc, block, style)
        elif block_type == "image":
            add_image_block(doc, block, style, base_dir)
        elif block_type == "heading":
            add_heading(doc, str(block.get("text", "")), int(block.get("level", 2)))
        elif block_type == "page_break":
            doc.add_page_break()
        elif block_type == "spacer":
            count = max(1, int(block.get("count", 1)))
            for _ in range(count):
                doc.add_paragraph()
        else:
            raise ValueError(f"Unsupported block type: {block_type}")


def add_executive_summary(doc: Document, summary: Mapping[str, Any], style: Mapping[str, Any]) -> None:
    if not summary:
        return
    add_heading(doc, str(summary.get("title", "执行摘要")), 1)
    overview = str(summary.get("overview", "") or "")
    if overview:
        add_paragraph_block(doc, {"type": "paragraph", "text": overview}, style)
    for label, key in (("关键决定", "decisions"), ("优先事项", "priorities"), ("主要风险", "risks")):
        items = summary.get(key) or []
        if items:
            p = doc.add_paragraph()
            add_text(p, label, style, bold=True, color=str(style["heading_color"]))
            add_list_block(doc, {"items": items}, style, numbered=False)


def add_sources(doc: Document, sources: Mapping[str, Any], style: Mapping[str, Any]) -> None:
    if not sources or not sources.get("entries"):
        return
    doc.add_page_break()
    add_heading(doc, str(sources.get("title", "来源与说明")), 1)
    for entry in sources.get("entries", []):
        if not isinstance(entry, Mapping):
            continue
        p = doc.add_paragraph()
        source_id = str(entry.get("id", ""))
        citation = str(entry.get("citation", ""))
        url = str(entry.get("url", "") or "")
        prefix = f"[{source_id}] " if source_id else ""
        add_text(p, prefix + citation, style, size=9.5)
        if url:
            add_text(p, f"  {url}", style, size=8.5, color="#666666")


def build(content_path: Path, output_path: Path, style_path: Optional[Path] = None) -> None:
    data = load_structured(content_path)
    style = deepcopy(DEFAULT_STYLE)
    if style_path:
        style = deep_merge(style, load_structured(style_path))
    style = deep_merge(style, data.get("style", {}) or {})
    meta = data.get("document", {}) or {}

    doc = Document()
    configure_document(doc, style, meta)
    if bool(style.get("cover_page", True)):
        add_cover(doc, meta, style)

    add_executive_summary(doc, data.get("executive_summary", {}) or {}, style)

    base_dir = content_path.parent
    for section in data.get("sections", []) or []:
        if not isinstance(section, Mapping):
            continue
        if bool(section.get("page_break_before", False)):
            doc.add_page_break()
        title = str(section.get("title", ""))
        if title:
            add_heading(doc, title, int(section.get("level", 1)))
        add_blocks(doc, section.get("blocks", []) or [], style, base_dir)

    add_sources(doc, data.get("sources", {}) or {}, style)

    appendices = data.get("appendices", []) or []
    if appendices:
        if bool(data.get("appendix_page_break", False)):
            doc.add_page_break()
        for appendix in appendices:
            if not isinstance(appendix, Mapping):
                continue
            if bool(appendix.get("page_break_before", False)):
                doc.add_page_break()
            add_heading(doc, str(appendix.get("title", "附录")), 1)
            add_blocks(doc, appendix.get("blocks", []) or [], style, base_dir)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a polished DOCX from YAML/JSON content.")
    parser.add_argument("content", type=Path, help="Path to report YAML or JSON")
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX path")
    parser.add_argument("--style", type=Path, help="Optional style YAML/JSON override")
    parser.add_argument("--force", action="store_true", help="Overwrite an existing output file")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if not args.content.exists():
        print(f"ERROR: content file not found: {args.content}", file=sys.stderr)
        return 2
    if args.style and not args.style.exists():
        print(f"ERROR: style file not found: {args.style}", file=sys.stderr)
        return 2
    if args.output.exists() and not args.force:
        print(f"ERROR: output already exists: {args.output}. Use --force to overwrite.", file=sys.stderr)
        return 2
    try:
        build(args.content.resolve(), args.output.resolve(), args.style.resolve() if args.style else None)
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
