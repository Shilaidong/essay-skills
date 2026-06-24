#!/usr/bin/env python3
"""Render a DOCX to PNG pages and produce a lightweight QA report.

This script prepares visual QA; it does not replace a human/agent inspection of
every page.  It prefers the canonical renderer from the host document skill and
falls back to LibreOffice + pdftoppm when available.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Sequence

try:
    from PIL import Image, ImageChops, ImageDraw, ImageFont, ImageStat
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency: Pillow. Run: pip install Pillow") from exc

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency: python-docx. Run: pip install python-docx") from exc


PLACEHOLDER_PATTERNS = {
    "TODO": re.compile(r"\bTODO\b", re.I),
    "TBD": re.compile(r"\bTBD\b", re.I),
    "待补充": re.compile(r"待补充|信息待补齐"),
    "XX placeholder": re.compile(r"(?<![A-Za-z])X{2,}(?![A-Za-z])", re.I),
    "tool citation token": re.compile(r"turn\d+(?:search|view|news|fetch|image)\d+", re.I),
    "sandbox URI": re.compile(r"sandbox:/", re.I),
    "reference marker": re.compile(r"\[\[REF:"),
}


def natural_key(path: Path) -> List[Any]:
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", path.name)]


def run_command(command: Sequence[str], *, env: Dict[str, str] | None = None) -> None:
    proc = subprocess.run(command, text=True, capture_output=True, env=env)
    if proc.returncode != 0:
        raise RuntimeError(
            f"Command failed ({proc.returncode}): {' '.join(command)}\n"
            f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
        )


def canonical_renderer() -> Path | None:
    candidates = [
        os.environ.get("DOCX_RENDERER"),
        "/home/oai/skills/docx/render_docx.py",
    ]
    for raw in candidates:
        if raw and Path(raw).exists():
            return Path(raw)
    return None


def render_with_canonical(docx_path: Path, render_dir: Path, emit_pdf: bool, verbose: bool) -> None:
    renderer = canonical_renderer()
    if renderer is None:
        raise FileNotFoundError("Canonical DOCX renderer not found")
    command = [sys.executable, str(renderer), str(docx_path), "--output_dir", str(render_dir)]
    if emit_pdf:
        command.append("--emit_pdf")
    if verbose:
        command.append("--verbose")
    run_command(command)


def render_with_fallback(docx_path: Path, render_dir: Path, emit_pdf: bool, verbose: bool) -> None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    pdftoppm = shutil.which("pdftoppm")
    if not soffice or not pdftoppm:
        raise RuntimeError("Fallback rendering needs LibreOffice/soffice and pdftoppm")

    render_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="docx-qa-") as temp_dir_raw:
        temp_dir = Path(temp_dir_raw)
        home = temp_dir / "home"
        profile = temp_dir / "profile"
        home.mkdir()
        profile.mkdir()
        env = os.environ.copy()
        env["HOME"] = str(home)
        command = [
            soffice,
            "--headless",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(temp_dir),
            str(docx_path),
        ]
        if verbose:
            print("Running:", " ".join(command))
        run_command(command, env=env)
        pdf_path = temp_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.exists():
            raise RuntimeError("LibreOffice did not produce a PDF")
        run_command([pdftoppm, "-png", "-r", "160", str(pdf_path), str(render_dir / "page")])
        if emit_pdf:
            shutil.copy2(pdf_path, render_dir / pdf_path.name)


def extract_text_and_structure(docx_path: Path) -> Dict[str, Any]:
    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs]
    table_text: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            table_text.append("\t".join(cell.text for cell in row.cells))
    full_text = "\n".join(paragraphs + table_text)

    image_count = 0
    with zipfile.ZipFile(docx_path) as zf:
        image_count = len([name for name in zf.namelist() if name.startswith("word/media/")])

    placeholders: Dict[str, int] = {}
    for label, pattern in PLACEHOLDER_PATTERNS.items():
        count = len(pattern.findall(full_text))
        if count:
            placeholders[label] = count

    return {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "embedded_image_count": image_count,
        "section_count": len(doc.sections),
        "character_count": len(full_text),
        "placeholder_hits": placeholders,
    }


def page_whiteness(image_path: Path) -> float:
    with Image.open(image_path).convert("RGB") as image:
        image.thumbnail((400, 400))
        stat = ImageStat.Stat(image)
        # Mean of normalized RGB; a value close to 1.0 is mostly white.
        return sum(stat.mean) / (3 * 255.0)


def make_contact_sheet(page_paths: List[Path], output: Path) -> None:
    if not page_paths:
        return
    thumbs: List[Image.Image] = []
    max_thumb_width = 360
    margin = 24
    label_height = 28
    for page in page_paths:
        image = Image.open(page).convert("RGB")
        ratio = max_thumb_width / image.width
        thumb = image.resize((max_thumb_width, int(image.height * ratio)))
        canvas = Image.new("RGB", (max_thumb_width, thumb.height + label_height), "white")
        canvas.paste(thumb, (0, label_height))
        draw = ImageDraw.Draw(canvas)
        draw.text((6, 6), page.stem, fill="black")
        thumbs.append(canvas)

    columns = 3 if len(thumbs) >= 3 else len(thumbs)
    rows = (len(thumbs) + columns - 1) // columns
    cell_width = max_thumb_width + margin
    cell_height = max(t.height for t in thumbs) + margin
    sheet = Image.new("RGB", (columns * cell_width + margin, rows * cell_height + margin), "white")
    for index, thumb in enumerate(thumbs):
        x = margin + (index % columns) * cell_width
        y = margin + (index // columns) * cell_height
        sheet.paste(thumb, (x, y))
    output.parent.mkdir(parents=True, exist_ok=True)
    sheet.save(output)
    for thumb in thumbs:
        thumb.close()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Render and inspect a DOCX for visual QA.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--render-dir", type=Path, required=True)
    parser.add_argument("--json-output", type=Path)
    parser.add_argument("--emit-pdf", action="store_true")
    parser.add_argument("--strict", action="store_true", help="Fail on placeholders or suspicious blank pages")
    parser.add_argument("--verbose", action="store_true")
    parser.add_argument("--force", action="store_true", help="Clear an existing render directory")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    docx_path = args.docx.resolve()
    render_dir = args.render_dir.resolve()
    if not docx_path.exists():
        print(f"ERROR: DOCX not found: {docx_path}", file=sys.stderr)
        return 2
    if render_dir.exists() and any(render_dir.iterdir()):
        if not args.force:
            print(f"ERROR: render directory is not empty: {render_dir}. Use --force.", file=sys.stderr)
            return 2
        shutil.rmtree(render_dir)
    render_dir.mkdir(parents=True, exist_ok=True)

    try:
        structure = extract_text_and_structure(docx_path)
        try:
            render_with_canonical(docx_path, render_dir, args.emit_pdf, args.verbose)
            renderer_used = "canonical"
        except Exception as canonical_error:  # noqa: BLE001
            if args.verbose:
                print(f"Canonical renderer unavailable/failed: {canonical_error}")
            render_with_fallback(docx_path, render_dir, args.emit_pdf, args.verbose)
            renderer_used = "libreoffice-fallback"

        pages = sorted(render_dir.glob("page-*.png"), key=natural_key)
        if not pages:
            # Fallback pdftoppm usually emits page-1.png; tolerate other naming.
            pages = sorted([p for p in render_dir.glob("*.png") if p.name != "contact-sheet.png"], key=natural_key)
        if not pages:
            raise RuntimeError("Renderer produced no page PNG files")

        blank_flags = []
        page_info = []
        for page in pages:
            with Image.open(page) as image:
                width, height = image.size
            whiteness = page_whiteness(page)
            suspicious_blank = whiteness > 0.9995
            if suspicious_blank:
                blank_flags.append(page.name)
            page_info.append({
                "file": page.name,
                "width": width,
                "height": height,
                "whiteness": round(whiteness, 4),
                "suspicious_blank": suspicious_blank,
            })

        contact_sheet = render_dir / "contact-sheet.png"
        make_contact_sheet(pages, contact_sheet)

        report = {
            "docx": str(docx_path),
            "renderer": renderer_used,
            "page_count": len(pages),
            "contact_sheet": str(contact_sheet),
            "structure": structure,
            "pages": page_info,
            "manual_review_required": True,
            "manual_review_checklist": [
                "Inspect every page at 100% zoom",
                "Check clipping, overlap, broken tables, missing glyphs and image resolution",
                "Check headings are not orphaned and page breaks are intentional",
                "Check citations, numbers, captions and page numbering",
            ],
        }

        json_path = args.json_output.resolve() if args.json_output else render_dir / "qa-report.json"
        json_path.parent.mkdir(parents=True, exist_ok=True)
        json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

        print(f"Rendered {len(pages)} page(s) to {render_dir}")
        print(f"Contact sheet: {contact_sheet}")
        print(f"QA report: {json_path}")
        if structure["placeholder_hits"]:
            print("Placeholder/tool-token hits:")
            for label, count in structure["placeholder_hits"].items():
                print(f"- {label}: {count}")
        if blank_flags:
            print("Suspiciously blank pages:")
            for name in blank_flags:
                print(f"- {name}")

        if args.strict and (structure["placeholder_hits"] or blank_flags):
            return 1
        return 0
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
